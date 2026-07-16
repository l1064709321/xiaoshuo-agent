"""文件读取与导出。

读取(供上传小说 → 分块入库):
  .txt .md .markdown .text     -> 直接 utf-8 解码
  .docx                        -> python-docx
  .pdf                         -> pypdf
  .epub                        -> ebooklib + bs4

导出(从项目章节生成可下载文件):
  .txt   -> 纯文本
  .md    -> Markdown
  .docx  -> Word (python-docx)
  .html  -> HTML (可浏览器打印为 PDF)

各解析库若缺失则优雅降级:返回错误说明,不崩溃服务。
"""
from __future__ import annotations

import io
import os
from typing import Optional


# ============ 读取 ============
def parse_bytes(filename: str, data: bytes) -> str:
    """根据文件名后缀解析为纯文本。失败抛 ValueError。"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md", ".markdown", ".text", ".csv", ".log"):
        return data.decode("utf-8", errors="ignore")
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".epub":
        return _parse_epub(data)
    # 未知后缀:尝试当文本
    return data.decode("utf-8", errors="ignore")


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise ValueError("读取 .docx 需要 python-docx,请 pip install python-docx") from e
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 顺便提取表格文本
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise ValueError("读取 .pdf 需要 pypdf,请 pip install pypdf") from e
    reader = PdfReader(io.BytesIO(data))
    texts = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t:
            texts.append(t)
    return "\n".join(texts)


def _parse_epub(data: bytes) -> str:
    try:
        import ebooklib  # type: ignore
        from ebooklib import epub  # type: ignore
    except ImportError as e:
        raise ValueError("读取 .epub 需要 ebooklib,请 pip install ebooklib") from e
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as e:
        raise ValueError("读取 .epub 需要 beautifulsoup4") from e
    book = epub.read_epub(io.BytesIO(data))
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        html = item.get_content().decode("utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        txt = soup.get_text(separator="\n").strip()
        if txt:
            parts.append(txt)
    return "\n\n".join(parts)


# ============ 导出 ============
def export_project(project: dict, chapters: list[dict], fmt: str) -> tuple[str, bytes, str]:
    """导出整个项目为指定格式。

    返回 (filename, content_bytes, content_type)。
    fmt: txt | md | docx | html
    """
    title = (project.get("name") or "novel").strip()
    safe = "".join(c for c in title if c.isalnum() or c in "._-") or "novel"
    ext_map = {"txt": "txt", "md": "md", "markdown": "md", "docx": "docx", "html": "html", "htm": "html"}
    ext = ext_map.get(fmt.lower(), "txt")

    if ext == "txt":
        text = _to_text(project, chapters)
        return f"{safe}.txt", text.encode("utf-8"), "text/plain; charset=utf-8"
    if ext == "md":
        text = _to_markdown(project, chapters)
        return f"{safe}.md", text.encode("utf-8"), "text/markdown; charset=utf-8"
    if ext == "docx":
        return f"{safe}.docx", _to_docx(project, chapters), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if ext == "html":
        html = _to_html(project, chapters)
        return f"{safe}.html", html.encode("utf-8"), "text/html; charset=utf-8"
    raise ValueError(f"不支持的导出格式: {fmt}")


def _meta_block(project: dict) -> str:
    lines = []
    if project.get("genre"):
        lines.append(f"类型:{project['genre']}")
    if project.get("style"):
        lines.append(f"文风:{project['style']}")
    if project.get("premise"):
        lines.append(f"核心设定:{project['premise']}")
    return "\n".join(lines)


def _sorted_chapters(chapters: list[dict]) -> list[dict]:
    return sorted(chapters, key=lambda c: c.get("idx", 0))


def _to_text(project: dict, chapters: list[dict]) -> str:
    out = [project.get("name", "未命名")]
    meta = _meta_block(project)
    if meta:
        out.append("")
        out.append(meta)
    out.append("")
    for ch in _sorted_chapters(chapters):
        out.append(f"第{ch.get('idx', 0) + 1}章 {ch.get('title', '')}")
        out.append("")
        out.append(ch.get("content") or "(本章尚未撰写)")
        out.append("")
        out.append("")
    return "\n".join(out)


def _to_markdown(project: dict, chapters: list[dict]) -> str:
    out = [f"# {project.get('name', '未命名')}"]
    meta = _meta_block(project)
    if meta:
        out.append("")
        out.append(f"> {meta.replace(chr(10), chr(10) + '> ')}")
    out.append("")
    out.append("---")
    out.append("")
    for ch in _sorted_chapters(chapters):
        out.append(f"## 第{ch.get('idx', 0) + 1}章 {ch.get('title', '')}")
        if ch.get("outline"):
            out.append("")
            out.append(f"*梗概:{ch['outline']}*")
        out.append("")
        out.append(ch.get("content") or "(本章尚未撰写)")
        out.append("")
    return "\n".join(out)


def _to_html(project: dict, chapters: list[dict]) -> str:
    import markdown as md  # type: ignore
    body = _to_markdown(project, chapters)
    html_body = md.markdown(body, extensions=["extra"])
    title = project.get("name", "未命名")
    return f"""<!DOCTYPE html>
<html lang="zh"><head><meta charset="utf-8">
<title>{title}</title>
<style>
body{{font-family:"PingFang SC","Microsoft YaHei",serif;max-width:760px;margin:40px auto;
  line-height:1.9;padding:0 20px;color:#222}}
h1{{text-align:center}} h2{{border-bottom:1px solid #ddd;padding-bottom:4px}}
@media print{{body{{margin:0;max-width:none}}}}
</style></head>
<body>{html_body}</body></html>"""


def _to_docx(project: dict, chapters: list[dict]) -> bytes:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError as e:
        raise ValueError("导出 .docx 需要 python-docx") from e
    doc = Document()
    # 标题
    h = doc.add_heading(project.get("name", "未命名"), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 元信息
    meta = _meta_block(project)
    if meta:
        p = doc.add_paragraph()
        for line in meta.split("\n"):
            r = p.add_run(line + "\n")
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph("—" * 20).alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 章节
    for ch in _sorted_chapters(chapters):
        doc.add_heading(f"第{ch.get('idx',0)+1}章 {ch.get('title','')}", level=1)
        if ch.get("outline"):
            po = doc.add_paragraph()
            ro = po.add_run(f"梗概:{ch['outline']}")
            ro.italic = True
            ro.font.size = Pt(10)
            ro.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        content = ch.get("content") or "(本章尚未撰写)"
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
